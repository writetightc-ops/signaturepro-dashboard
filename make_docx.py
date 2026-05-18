"""
Генерирует ИНСТРУКЦИЯ.docx из ИНСТРУКЦИЯ.md
Запуск: python make_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, os

MD = os.path.join(os.path.dirname(__file__), "ИНСТРУКЦИЯ.md")
OUT = os.path.join(os.path.dirname(__file__), "ИНСТРУКЦИЯ.docx")

# ─── цвета ───────────────────────────────────────────────────────────────────
C_TITLE   = RGBColor(0x1A, 0x1A, 0x2E)   # тёмно-синий заголовок
C_H1      = RGBColor(0x16, 0x21, 0x3E)
C_H2      = RGBColor(0x0F, 0x3C, 0x78)
C_H3      = RGBColor(0x1A, 0x6B, 0xAB)
C_H4      = RGBColor(0x2E, 0x86, 0xC1)
C_TH_BG   = RGBColor(0x0F, 0x3C, 0x78)   # шапка таблицы
C_ALT_BG  = RGBColor(0xE8, 0xF1, 0xFB)   # чётные строки
C_NOTE_BG = RGBColor(0xFF, 0xF9, 0xC4)   # цитаты (жёлтый)
C_NOTE_BD = RGBColor(0xF5, 0xA6, 0x23)   # бордер цитат
C_WARN    = RGBColor(0xFF, 0xE0, 0x82)   # ⚠️ предупреждение


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def set_cell_border(cell, sides=("top", "bottom", "left", "right"),
                    color="BBBBBB", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def bold_runs(para, text: str):
    """Добавляет текст в параграф, поддерживая **bold** разметку."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


def strip_bold(text: str) -> str:
    return re.sub(r"\*\*([^*]+)\*\*", r"\1", text)


def strip_inline(text: str) -> str:
    """Убирает `code`, **bold**, markdown-эмодзи-флаги."""
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    return text


def add_heading(doc: Document, text: str, level: int):
    colors = {1: C_H1, 2: C_H2, 3: C_H3, 4: C_H4}
    sizes  = {1: 18,   2: 14,   3: 12,   4: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after  = Pt(4)
    if level == 1:
        p.paragraph_format.left_indent = Cm(0)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = colors.get(level, C_H4)
    return p


def add_note(doc: Document, text: str, warn=False):
    """Блок > цитаты."""
    text = strip_inline(text.lstrip("> ").lstrip("⚠️").strip())
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    # левая граница
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:color"), f"{C_WARN[0]:02X}{C_WARN[1]:02X}{C_WARN[2]:02X}" if warn
             else f"{C_NOTE_BD[0]:02X}{C_NOTE_BD[1]:02X}{C_NOTE_BD[2]:02X}")
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(("⚠  " if warn else "ℹ  ") + text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x5D, 0x4E, 0x37) if warn else RGBColor(0x2C, 0x5F, 0x8A)


def add_table_from_rows(doc: Document, rows: list[list[str]]):
    if not rows or len(rows) < 2:
        return
    col_n = len(rows[0])
    t = doc.add_table(rows=len(rows), cols=col_n)
    t.style = "Table Grid"

    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = t.cell(ri, ci)
            cell_text = strip_inline(cell_text.strip())
            p = cell.paragraphs[0]
            p.clear()
            bold_runs(p, cell_text)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            for run in p.runs:
                run.font.size = Pt(9.5)

            if ri == 0:
                # шапка
                set_cell_bg(cell, C_TH_BG)
                set_cell_border(cell, color="FFFFFF", sz="6")
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                bg = C_ALT_BG if ri % 2 == 0 else RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_bg(cell, bg)
                set_cell_border(cell, color="CCCCCC", sz="4")


def parse_md_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        if re.match(r"\s*\|[-: |]+\|\s*$", line):
            continue
        parts = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(parts)
    return rows


def build_doc():
    doc = Document()

    # страница
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # базовый шрифт
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    with open(MD, encoding="utf-8") as f:
        lines = f.read().splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]

        # ── горизонтальный разделитель ───────────────────────────────────────
        if re.match(r"^---+$", line.strip()):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "6")
            bot.set(qn("w:color"), "AAAAAA")
            pBdr.append(bot)
            pPr.append(pBdr)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            i += 1
            continue

        # ── заголовки ────────────────────────────────────────────────────────
        m = re.match(r"^(#{1,4})\s+(.+)", line)
        if m:
            level = len(m.group(1))
            text  = strip_inline(m.group(2))
            # Для h1 (только один — главный заголовок) отдельный стиль
            if level == 1:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(10)
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(22)
                run.font.color.rgb = C_TITLE
            else:
                add_heading(doc, text, level)
            i += 1
            continue

        # ── таблица ──────────────────────────────────────────────────────────
        if line.startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            rows = parse_md_table(tbl_lines)
            add_table_from_rows(doc, rows)
            doc.add_paragraph()  # отступ после таблицы
            continue

        # ── цитаты / примечания ──────────────────────────────────────────────
        if line.startswith(">"):
            warn = "⚠" in line
            add_note(doc, line, warn=warn)
            i += 1
            continue

        # ── нумерованный список ──────────────────────────────────────────────
        m = re.match(r"^(\d+)\.\s+(.+)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent   = Cm(0.6)
            p.paragraph_format.space_before  = Pt(2)
            p.paragraph_format.space_after   = Pt(2)
            bold_runs(p, strip_inline(m.group(2)))
            for run in p.runs:
                run.font.size = Pt(10.5)
            i += 1
            continue

        # ── маркированный список ─────────────────────────────────────────────
        m = re.match(r"^[-*]\s+(.+)", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent  = Cm(0.6)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            bold_runs(p, strip_inline(m.group(1)))
            for run in p.runs:
                run.font.size = Pt(10.5)
            i += 1
            continue

        # ── пустая строка ────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── обычный параграф ─────────────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        bold_runs(p, strip_inline(line))
        for run in p.runs:
            run.font.size = Pt(10.5)
        i += 1

    doc.save(OUT)
    print(f"Готово: {OUT}")


if __name__ == "__main__":
    build_doc()
