from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Поля страницы ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Вспомогательные функции ────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="DDDDDD"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def para(text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # тёмный фон через shading абзаца
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1A1A1A')
    pPr.append(shd)
    # отступ внутри
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '200')
    pPr.append(ind)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p

def block_num(num, title, subtitle=""):
    """Заголовок блока с номером"""
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # номер
    c0 = t.cell(0, 0)
    c0.width = Cm(1.5)
    set_cell_bg(c0, '1A1A1A')
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(str(num))
    r0.bold = True
    r0.font.size = Pt(16)
    r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # заголовок
    c1 = t.cell(0, 1)
    set_cell_bg(c1, 'FAFAF9')
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(14)
    if subtitle:
        p1.add_run('\n')
        rs = p1.add_run(subtitle)
        rs.font.size = Pt(10)
        rs.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def copy_box(label, text, is_good=True):
    """Зелёная / красная коробка"""
    color = 'F0FDF4' if is_good else 'FFF5F5'
    border = '86EFAC' if is_good else 'FECACA'
    label_color = '16a34a' if is_good else 'DC2626'
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.cell(0, 0)
    set_cell_bg(c, color)
    set_cell_borders(c, border.replace('#',''))
    p = c.paragraphs[0]
    r_label = p.add_run(label + '\n')
    r_label.bold = True
    r_label.font.size = Pt(9)
    r_label.font.color.rgb = RGBColor(*bytes.fromhex(label_color))
    r_text = p.add_run(text)
    r_text.font.size = Pt(11)
    r_text.italic = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def copy_block(label, main_text, sub_text=""):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.cell(0, 0)
    set_cell_bg(c, 'F8F8F6')
    set_cell_borders(c, '1A1A1A')
    p = c.paragraphs[0]
    r_label = p.add_run(label.upper() + '\n')
    r_label.font.size = Pt(8)
    r_label.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    r_main = p.add_run(main_text)
    r_main.bold = True
    r_main.font.size = Pt(13)
    if sub_text:
        r_sub = p.add_run('\n' + sub_text)
        r_sub.font.size = Pt(10)
        r_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def note_box(title, text, bg='FFFBEB', border='FDE68A', title_color='92400E', text_color='78350F'):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.cell(0, 0)
    set_cell_bg(c, bg)
    set_cell_borders(c, border)
    p = c.paragraphs[0]
    r_t = p.add_run('⚠ ' + title.upper() + '\n')
    r_t.bold = True
    r_t.font.size = Pt(9)
    r_t.font.color.rgb = RGBColor(*bytes.fromhex(title_color))
    r_text = p.add_run(text)
    r_text.font.size = Pt(10)
    r_text.font.color.rgb = RGBColor(*bytes.fromhex(text_color))
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def info_box(title, text):
    note_box(title, text, bg='EFF6FF', border='BFDBFE', title_color='1E40AF', text_color='1E3A8A')

def bullet(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5) if indent else Cm(0)
    run = p.add_run('— ' + text)
    run.font.size = Pt(11)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E5E5E5')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_pricing_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # заголовки
    for i, h in enumerate(headers):
        c = t.cell(0, i)
        set_cell_bg(c, '1A1A1A')
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if col_widths:
            c.width = Cm(col_widths[i])
    # строки
    for ri, row in enumerate(rows):
        bg = 'FAFAF9' if ri % 2 == 1 else 'FFFFFF'
        for ci, cell_text in enumerate(row):
            c = t.cell(ri + 1, ci)
            set_cell_bg(c, bg)
            p = c.paragraphs[0]
            r = p.add_run(str(cell_text))
            r.font.size = Pt(10)
            if col_widths:
                c.width = Cm(col_widths[ci])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ══════════════════════════════════════════════════════════════
#  ТИТУЛЬНАЯ СТРАНИЦА
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ТЗ: Новый лендинг SignaturePro')
r.bold = True
r.font.size = Pt(28)

para('Техническое задание для веб-дизайнера', size=14, color='888888', align=WD_ALIGN_PARAGRAPH.CENTER)
para('Отдельная посадочная страница для тестирования нового позиционирования', size=12, color='AAAAAA', align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph().paragraph_format.space_after = Pt(30)

# Мета-таблица
t = doc.add_table(rows=5, cols=2)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
meta = [
    ('Проект', 'SignaturePro / Write Tight'),
    ('Тип страницы', 'Отдельный лендинг (не главная)'),
    ('Цель', 'Трафик на новое позиционирование'),
    ('Язык', 'Английский (США)'),
    ('Оригинал', 'signatureprodesign.com'),
]
for i, (k, v) in enumerate(meta):
    c0, c1 = t.cell(i, 0), t.cell(i, 1)
    set_cell_bg(c0, 'F0F0EF')
    r0 = c0.paragraphs[0].add_run(k)
    r0.bold = True
    r0.font.size = Pt(10)
    r1 = c1.paragraphs[0].add_run(v)
    r1.font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  ВВЕДЕНИЕ
# ══════════════════════════════════════════════════════════════
heading1('О документе')
para('Новая страница — самостоятельный лендинг с обновлённым позиционированием. Структура не повторяет главный сайт, но сохраняет ключевые блоки. Цель — протестировать новые смыслы на платном трафике.', size=11)

heading2('Что меняется в позиционировании')
for item in [
    'Уходим от «подписи, которой восхищаются» → инструмент для профессионала',
    'Акцент на деловую аудиторию: менеджмент, предприниматели',
    'Вводим концепт Stress Signature как уникальное УТП',
    'Убираем язык скидок и инфобиза',
    'Тариф «Премьер» — подпись под ваш почерк (анализ рукописного текста)',
]:
    bullet(item)

heading2('Тон коммуникации')
for item in [
    'Сдержанно и уверенно — без восторженности',
    'Прагматично — через ситуации, а не через эмоции',
    'Без «ограниченное предложение», таймеров, «успей»',
    'Отзывы — фактические, не восторженные',
    'Ориентир: «уверенный специалист», не инфобиз',
]:
    bullet(item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  БЛОК 1 — HERO
# ══════════════════════════════════════════════════════════════
block_num(1, 'Главный экран (Hero)', 'Первое, что видит пользователь. Ответить: «что это и для кого» за 5 секунд.')

copy_box('❌ БЫЛО (текущий сайт)', '«Express your personality with an exclusive signature crafted by pro calligraphers» — акцент на личность, эксклюзивность, самовыражение.', is_good=False)
copy_box('✅ СТАЛО', 'Акцент: профессиональный инструмент, соответствие уровню карьеры, функциональность.', is_good=True)

heading2('Копирайтинг — главный экран')

copy_block('Основной заголовок (H1)', 'You\'ve outgrown your signature.', 'Провокация через ситуацию. Говорим: ты уже вырос, подпись должна соответствовать.')
copy_block('Подзаголовок', 'A handcrafted signature designed to perform — in daily documents, high-stakes contracts, and everything in between.', '«Perform» — слово из бизнес-языка, не из языка красоты.')
copy_block('Кнопка CTA (основная)', 'Get My Signature')
copy_block('Кнопка CTA (вторичная)', 'See Examples →')
copy_block('Под кнопками — мелкий текст', '20,500+ professionals trust their signature to us · Since 2017 · 100% handcrafted')

note_box('Комментарий дизайнеру', 'Фон — тёмный или нейтральный. Никаких ярких акцентных цветов. Никаких баннеров со скидками на герое. Визуал: процесс написания, рука с пером, готовая подпись — сдержанно. Типографика крупная, серьёзный бизнес-тон.')

divider()

# ══════════════════════════════════════════════════════════════
#  БЛОК 2 — РЕГАЛИИ
# ══════════════════════════════════════════════════════════════
block_num(2, 'Регалии и партнёры', 'Быстрое подтверждение надёжности. Сразу после Hero.')

heading2('Что показываем')
for item in [
    '[Обязательно] Логотип DocuSign — партнёр, электронные подписи в серьёзном контексте',
    '[Обязательно] Цифры: 20,500+ клиентов · 8+ лет · 100% handcrafted (no AI)',
    '[Опционально] Другие логотипы партнёров или медиа-упоминаний',
    '[Опционально] Значок «No AI generation» — важный отличитель от конкурентов',
]:
    bullet(item)

copy_block('Заголовок блока (небольшой, серый)', 'Trusted by professionals worldwide')

note_box('Комментарий дизайнеру', 'Логотипы в серых тонах (не цветные) — стандарт для серьёзных сайтов. Цифры крупные. Блок минимальный, не должен кричать — просто говорит факты.')

divider()

# ══════════════════════════════════════════════════════════════
#  БЛОК 3 — АУДИТОРИЯ
# ══════════════════════════════════════════════════════════════
block_num(3, 'Целевая аудитория', 'Пользователь должен узнать себя. Описываем через ситуации, не через профессии.')

info_box('Ключевой принцип', 'Не писать «для юристов и финансистов». Описывать через ситуации: «когда вы подписываете контракт перед клиентом». Клиент сам узнаёт себя.')

copy_box('❌ БЫЛО', 'Никакой явной аудитории. Размытый оффер — «для всех».', is_good=False)
copy_box('✅ СТАЛО', 'Четыре категории через конкретные ситуации. Менеджмент — на первом месте.', is_good=True)

copy_block('Заголовок блока (H2)', 'A signature that works as hard as you do.', 'Whether you sign documents daily or seal major deals — your signature is the last impression you leave.')

heading2('4 карточки аудитории')

add_pricing_table(
    ['Карточка', 'Заголовок', 'Подтекст карточки'],
    [
        ['1 — Менеджмент', 'Executives & Managers', 'When your signature appears on contracts, approvals and agreements — it should reflect where you are now, not where you started.'],
        ['2 — Предприниматели', 'Business Owners', 'You\'ve built something real. Your signature is often the only handwritten detail in an otherwise digital business — make it count.'],
        ['3 — Специалисты', 'Licensed Professionals', 'Lawyers, doctors, consultants — your signature is a legal mark and a professional signal.'],
        ['4 — Карьерный старт', 'Career Milestone', 'New position. New company. New chapter. The moment you step up — your signature should too.'],
    ],
    col_widths=[3.5, 4, 8.5]
)

note_box('Комментарий дизайнеру', 'Карточки — сдержанные, без ярких цветов. Иконки — минималистичные (линейные). Никаких фотографий людей. Менеджмент — первая карточка.')

divider()

# ══════════════════════════════════════════════════════════════
#  БЛОК 4 — ПРИМЕРЫ РАБОТ
# ══════════════════════════════════════════════════════════════
block_num(4, 'Примеры работ', 'Портфолио без имён клиентов, но с профессиональным контекстом.')

info_box('Ключевой принцип', 'Это главное социальное доказательство на сайте. 8–12 подписей. К каждой — короткий контекст профессии без имени клиента.')

copy_box('❌ БЫЛО', '«Our works» — галерея без контекста. Широкий микс профессий включая фармацевтов.', is_good=False)
copy_box('✅ СТАЛО', 'Галерея с контекстом профессии. Менеджмент и предприниматели — в приоритете. Фармацевты убраны или в конец.', is_good=True)

copy_block('Заголовок блока', 'Our work speaks for itself.', 'A selection of signatures we\'ve designed for professionals across industries. Client names are kept private.')

heading2('Подписи к примерам работ')
for item in [
    'Signature for a Finance Director',
    'Signature for a Real Estate Executive',
    'Signature for a Tech Startup Founder',
    'Signature for a Corporate Lawyer',
    'Signature for a General Manager',
    'Signature for a Consulting Partner',
    'Signature for a Sales Director',
    'Signature for a Physician (Private Practice)',
]:
    bullet(item)

note_box('Комментарий дизайнеру', 'Сетка или горизонтальный скролл. Подписи на белом фоне. Подпись к изображению — мелкая серая. 2–3 варианта «до/после» без акцента на «было плохо».')

divider()

# ══════════════════════════════════════════════════════════════
#  БЛОК 5 — STRESS SIGNATURE
# ══════════════════════════════════════════════════════════════
block_num(5, 'Описание продукта — Stress Signature', 'Уникальное УТП, которого нет ни у одного конкурента в мире.')

info_box('Почему это важно', 'Ни один конкурент не говорит о разрыве между подписью в идеальных условиях и в реальности. Это главное отличие Write Tight от всего рынка.')

copy_block('Заголовок блока', 'A signature that holds up under pressure.', 'Most signatures look great at a desk. The problem is the rest of the time.')

heading2('Объяснение концепта')
copy_block('Подзаголовок секции', 'Introducing Stress Signature', 'We design two versions of your signature.\n\nYour full signature — crafted for formal documents and moments that matter.\n\nYour Stress Signature — a streamlined version designed for real conditions: quick approvals, tablet signing, the end of a long day. Same character, built to perform when you\'re not at your best.')

heading2('3 ситуации (плашки)')
add_pricing_table(
    ['Ситуация', 'Текст'],
    [
        ['Подписание на планшете', 'No surface tension, awkward angle — your stress signature handles it.'],
        ['40 согласований в спешке', 'Speed without falling apart — designed to look intentional, not hurried.'],
        ['Важный контракт', 'When eyes are on you — your full signature delivers the impression it should.'],
    ],
    col_widths=[5, 11]
)

note_box('Комментарий дизайнеру', 'Ключевой блок страницы. Визуально показать рядом два варианта подписи: «Full» и «Stress». Stress Signature включён в Optimal и Premier — не в Standard. Это важно показать визуально.')

divider()

# ══════════════════════════════════════════════════════════════
#  БЛОК 6 — ПРОЦЕСС
# ══════════════════════════════════════════════════════════════
block_num(6, 'Процесс работы', 'Снимаем тревогу: покупатель понимает, что именно произойдёт после оплаты.')

copy_block('Заголовок блока', 'How it works', 'Simple. Exactly what it sounds like.')

heading2('4 шага процесса')
add_pricing_table(
    ['Шаг', 'Название', 'Описание', 'Срок'],
    [
        ['01', 'Fill in a brief', 'Tell us about your preferences, style, and how you\'ll use your signature. Takes less than 5 minutes.', '5 мин'],
        ['02', 'We design options', 'Our calligrapher creates your signature options (3 for Standard & Optimal, 1 tailored variant for Premier).', '3 дня'],
        ['03', 'You choose & refine', 'Pick your favorite, or request edits. 3 rounds of revisions included in every plan.', 'По запросу'],
        ['04', 'Learn to sign it', 'Your calligrapher records a personal video tutorial — just for your signature. 30 minutes to master it.', '24 часа'],
    ],
    col_widths=[1.5, 4, 9, 2]
)

note_box('Комментарий дизайнеру', 'Шаги — горизонтальные, пронумерованные. Минималистично. Между шагами — стрелки или линии. Добавить плашку: «Your manager will reach out via email within 24 hours of your order».')

divider()

# ══════════════════════════════════════════════════════════════
#  БЛОК 7 — ТАРИФЫ
# ══════════════════════════════════════════════════════════════
block_num(7, 'Тарифы', 'Три чётких плана. Без скидок. Без «limited time». Без инфобиза.')

copy_box('❌ БЫЛО', '$69 / $140 / $269 со скидками 50%. Акции, баннеры, «успей». Без чёткой разницы в ценности.', is_good=False)
copy_box('✅ СТАЛО', '$99 / $159 / $249 без скидок. Optimal — рекомендованный (Stress Signature включён). Premier — анализ почерка.', is_good=True)

copy_block('Заголовок блока', 'Choose your plan', 'Every plan includes a personal video tutorial and our handwriting practice course. No subscriptions, no hidden fees.')

heading2('Тарифная таблица')
add_pricing_table(
    ['Параметр', 'Standard · $99', 'Optimal · $159 ⭐', 'Premier · $249'],
    [
        ['Варианты подписи', '3 варианта', '3 варианта', '1 вариант (под ваш почерк)'],
        ['Правки', '✓ 3 правки', '✓ 3 правки', '✓ 3 правки'],
        ['Stress Signature', '✗ Нет', '✓ Включено', '✓ Включено'],
        ['Видеообучение', '✓ Да', '✓ Да', '✓ Да'],
        ['Курс Learn to Write. Again', '✓ Да', '✓ Да', '✓ Да'],
        ['Анализ почерка', '✗ Нет', '✗ Нет', '✓ Рукописный анализ'],
        ['Кнопка CTA', 'Get Standard', 'Get Optimal', 'Get Premier'],
    ],
    col_widths=[5, 4, 4.5, 4.5]
)

heading2('Описания карточек')
copy_block('Standard — $99', 'The foundation. A professionally designed signature with everything you need to start using it immediately.')
copy_block('Optimal — $159 (рекомендуемый)', 'The complete package. Includes Stress Signature — a streamlined version designed for real-world signing conditions.')
copy_block('Premier — $249', 'Fully personalized. We analyze your natural handwriting and craft a signature that feels like an evolved version of you — not a replacement.')

note_box('Комментарий дизайнеру', 'Карточка Optimal — выделена тёмным фоном. «Most popular» — маленькая пометка сверху. Никаких «СКИДКА 50%» или таймеров. Под тарифами — мелко: «You can pay with PayPal, Apple Pay, Visa, Mastercard.»')

divider()

# ══════════════════════════════════════════════════════════════
#  БЛОК 8 — PREMIER DEEP DIVE
# ══════════════════════════════════════════════════════════════
block_num(8, 'Тариф «Premier» — детальный блок', 'Объяснение уникальной механики анализа почерка.')

info_box('Зачем отдельный блок', '«Подпись под ваш почерк» — сложная концепция. Клиент должен понять, чем Premier принципиально отличается от Optimal. Этот блок продаёт самый дорогой тариф.')

copy_block('Заголовок блока', 'Premier: built around your handwriting.', 'Not a signature designed for you. A signature evolved from you.')

copy_block('Основной текст (как это работает)', 'How Premier is different', 'In Standard and Optimal, our calligrapher designs a signature based on your preferences.\n\nIn Premier, we start with your actual handwriting.\n\nYou write a few words — your name, a phrase, anything natural. Our calligrapher analyzes how your hand moves: the pressure, the rhythm, the natural letter connections that are yours alone.\n\nThe result is a signature that feels like something you\'ve always had — refined, not replaced.')

copy_block('Аналогия (помогает продать разницу)', 'Like a bespoke suit — cut to your exact measurements, not just your size.')

heading2('3 шага процесса Premier')
for item in [
    'You write several words in your natural handwriting — no preparation needed',
    'Our calligrapher analyzes the rhythm, pressure, and natural letter shapes in your writing',
    'We craft one signature that builds on your existing hand — not a template, not a catalogue style',
]:
    bullet(item)

note_box('Комментарий дизайнеру', 'Визуально: 3-шаговая схема — «рукописный текст» → «анализ» → «готовая подпись». Блок должен визуально отличаться от остальных — другой фон или рамка. Он продаёт самый дорогой тариф.')

divider()

# ══════════════════════════════════════════════════════════════
#  БЛОК 9 — ОТЗЫВЫ
# ══════════════════════════════════════════════════════════════
block_num(9, 'Отзывы', 'Сдержанные и фактические. Без восклицательных знаков и «я в восторге!!!»')

copy_box('❌ БЫЛО', '«OMG I\'m so obsessed!!! Best purchase of my life!!!» — слишком эмоционально для деловой аудитории.', is_good=False)
copy_box('✅ СТАЛО', 'Фактические отзывы с профессиональным контекстом. Не восторг — профессиональная оценка.', is_good=True)

copy_block('Заголовок блока', 'What our clients say')

heading2('Формат отзывов — примеры')
add_pricing_table(
    ['Имя / контекст', 'Текст отзыва'],
    [
        ['Tom G. · Real Estate, Florida', '«I ordered before taking on a new partner. The process was straightforward, and I\'ve been using the signature for 8 months. It holds up exactly as described — including on tablet signings.»'],
        ['M. Brown · Photography, New York', '«Got Premier. The result was closer to my natural writing than I expected. Took about a week to get comfortable, now it\'s automatic.»'],
        ['A. Creasey · Business Coach', '«I\'d put off doing this for two years. Took 30 minutes to learn. Worth it.»'],
    ],
    col_widths=[4.5, 11.5]
)

heading2('Обязательные требования к отзывам')
for item in [
    '[Обязательно] Имя + профессия/сфера под каждым отзывом',
    '[Обязательно] Убрать восклицательные знаки при редактировании',
    '[Опционально] Фото клиента (если дали разрешение)',
    '[Опционально] Instagram @handle — кликабельный',
]:
    bullet(item)

note_box('Комментарий дизайнеру', 'Карточки на светлом фоне, без звёздочек. Формат «цитата + имя + должность». Видеоотзывы — отдельная полоса с тёмным фоном.')

divider()

# ══════════════════════════════════════════════════════════════
#  БЛОК 10 — FAQ
# ══════════════════════════════════════════════════════════════
block_num(10, 'Часто задаваемые вопросы (FAQ)', 'Снимаем последние возражения перед покупкой.')

copy_block('Заголовок блока', 'Questions')

heading2('Вопросы и ответы')
add_pricing_table(
    ['Вопрос', 'Ответ'],
    [
        ['My handwriting is really bad. Is it a problem?', 'Not at all. We design from scratch — or in Premier, we take what\'s there and refine it.'],
        ['What exactly is Stress Signature?', 'A simplified version of your main signature — designed to look intentional when signing quickly, on a tablet, or under pressure. Included in Optimal and Premier.'],
        ['Will my new signature be legally binding?', 'Yes. A legal signature only requires consistent use and intent. Your new signature is exactly as valid as any other.'],
        ['How do I learn to sign it?', 'Your calligrapher records a personal video tutorial filmed specifically for your signature. Most clients feel confident in 30 minutes.'],
        ['What\'s in the «Learn to Write. Again» course?', 'A 7-day program to improve your overall handwriting. Created by Ann Storm, our chief calligrapher. Included in every plan.'],
        ['How long does it take?', 'First options within 3 business days. Tutorial delivered within 24 hours after approval.'],
        ['What if I\'m not satisfied?', '3 rounds of revisions included. If still not satisfied — we offer a refund.'],
        ['Do you share my signature publicly?', 'Only with your explicit permission. You can opt out entirely.'],
    ],
    col_widths=[6, 10]
)

note_box('Комментарий дизайнеру', 'Формат аккордеона (раскрывающиеся вопросы). Под FAQ — финальный CTA.')

divider()

# ══════════════════════════════════════════════════════════════
#  БЛОК 11 — ПОДВАЛ
# ══════════════════════════════════════════════════════════════
block_num(11, 'Подвал сайта (Footer)', 'Финальный CTA + контакты + юридика.')

copy_block('Финальный CTA перед подвалом', 'Your signature is the only handwritten part of your professional life.', 'Make it intentional.\n\n[Кнопка]: Get My Signature')

heading2('Состав подвала')
for item in [
    '[Обязательно] Логотип SignaturePro',
    '[Обязательно] Email поддержки',
    '[Обязательно] Телефон: +1 (754) 329-3011',
    '[Обязательно] Адрес: 1849 S Ocean Dr, Hallandale Beach, FL 33009',
    '[Обязательно] Часы работы: Mon–Fri 10AM–6PM, Sat 12AM–6PM, Sun closed',
    '[Обязательно] Ссылки: Privacy Policy · Refund Policy · Terms of Service',
    '[Обязательно] © 2017–2025 USA CHECK NY INC. All rights reserved.',
    '[Опционально] Ссылка на Instagram',
    '[Опционально] Ссылка на главный сайт signatureprodesign.com',
]:
    bullet(item)

note_box('Комментарий дизайнеру', 'Тёмный подвал. Лаконично. Никаких дополнительных блоков с акциями. Перед подвалом — полноширинный CTA-блок с одним заголовком и одной кнопкой.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  СВОДНАЯ ТАБЛИЦА ИЗМЕНЕНИЙ
# ══════════════════════════════════════════════════════════════
heading1('Сводная таблица изменений')
para('Что убрать, что добавить, что изменить относительно текущего сайта signatureprodesign.com', size=11, color='888888')

add_pricing_table(
    ['Элемент', 'БЫЛО', 'СТАЛО'],
    [
        ['Главный слоган', '«Express your personality»', '«You\'ve outgrown your signature.»'],
        ['Целевая аудитория', 'Все, кто хочет красивую подпись', 'Менеджмент, предприниматели, специалисты'],
        ['Цены', '$69 / $140 / $269 со скидками 50%', '$99 / $159 / $249 без скидок'],
        ['Ключевое УТП', 'Handcrafted, exclusive', 'Stress Signature — работает в реальных условиях'],
        ['Тариф Premier', 'Не объяснён, нет механики', 'Анализ почерка — «как костюм по меркам»'],
        ['Отзывы', 'Восторженные, эмоциональные', 'Фактические, с профессиональным контекстом'],
        ['Акции и таймеры', '«50% off», таймеры, «limited offer»', 'Убрать полностью'],
        ['Курс в тарифе', 'Отдельный акцент', 'Включён во все тарифы, не акцентируется'],
        ['Портфолио', 'Широкий микс, включая фармацевтов', 'Менеджмент и бизнес — в приоритете'],
    ],
    col_widths=[4.5, 6, 6]
)

# ── Сохранение ─────────────────────────────────────────────────
path = r'd:\Подписи курсор\ТЗ_лендинг_SignaturePro.docx'
doc.save(path)
print(f'Сохранено: {path}')
